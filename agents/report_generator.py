"""
Phase 4 — Report Generator

Produces a professional Word (.docx) report summarising the financial analysis.
Structure:
  1. Cover Page
  2. Executive Summary
  3. Financial Performance — Revenue & Profitability
  4. Balance Sheet Health
  5. Cash Flow Analysis
  6. Key Ratios Dashboard (table)
  7. Valuation Summary
  8. Key Risks & Strengths
  9. Disclaimer
"""
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL

from config.settings import COLOR_GOLD, COLOR_NAVY, REPORTS_DIR
from models.company_data import FinancialData
from utils.logger import get_logger

log = get_logger()

# Colour helpers
_NAVY  = RGBColor(0x1F, 0x38, 0x64)
_GOLD  = RGBColor(0xC9, 0xA8, 0x4C)
_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
_LGRAY = RGBColor(0xF2, 0xF2, 0xF2)
_RED   = RGBColor(0xC0, 0x00, 0x00)
_GREEN = RGBColor(0x00, 0x60, 0x00)


_WML_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _set_cell_bg(cell, hex_color: str):
    """Set table cell background colour."""
    from docx.oxml import parse_xml
    shading = parse_xml(
        f'<w:shd xmlns:w="{_WML_NS}" '
        f'w:val="clear" w:color="auto" w:fill="{hex_color.lstrip("#")}"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def _para_style(para, size: int = 11, bold: bool = False,
                color: RGBColor = None, align: str = "left", italic: bool = False):
    para.alignment = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }.get(align, WD_ALIGN_PARAGRAPH.LEFT)
    for run in para.runs:
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic
        if color:
            run.font.color.rgb = color


def _heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = _NAVY
        run.font.bold = True
    return p


def _body(doc: Document, text: str):
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.size = Pt(11)
    return p


def _fmt_num(value: Optional[float], decimals: int = 0, pct: bool = False) -> str:
    if value is None:
        return "—"
    if pct:
        return f"{value * 100:.{decimals}f}%"
    if abs(value) >= 1_000:
        return f"{value:,.{decimals}f}"
    return f"{value:.{decimals}f}"


def _add_kv_table(doc: Document, rows: List[tuple], col_widths=(3.0, 1.8)):
    """Add a two-column label/value table."""
    table = doc.add_table(rows=len(rows) + 1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr = table.rows[0]
    for cell, text, color in zip(hdr.cells, ["Metric", "Value"], [_NAVY, _NAVY]):
        cell.text = text
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = _WHITE
        _set_cell_bg(cell, "1F3864")
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for i, (label, value) in enumerate(rows):
        row = table.rows[i + 1]
        row.cells[0].text = label
        row.cells[1].text = str(value)
        row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if i % 2 == 0:
            _set_cell_bg(row.cells[0], "EBF3FB")
            _set_cell_bg(row.cells[1], "EBF3FB")

    # Column widths
    for row in table.rows:
        row.cells[0].width = Inches(col_widths[0])
        row.cells[1].width = Inches(col_widths[1])

    doc.add_paragraph()
    return table


def _add_multi_col_table(doc: Document, headers: List[str], rows: List[List[str]]):
    """Add a multi-column data table."""
    n_cols = len(headers)
    table = doc.add_table(rows=len(rows) + 1, cols=n_cols)
    table.style = "Table Grid"

    hdr = table.rows[0]
    for cell, text in zip(hdr.cells, headers):
        cell.text = text
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = _WHITE
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_cell_bg(cell, "1F3864")

    for i, row_data in enumerate(rows):
        row = table.rows[i + 1]
        for j, val in enumerate(row_data):
            row.cells[j].text = str(val)
            if j > 0:
                row.cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            if i % 2 == 0:
                _set_cell_bg(row.cells[j], "EBF3FB")

    doc.add_paragraph()
    return table


# ══════════════════════════════════════════════════════════════════════════════
# Section builders
# ══════════════════════════════════════════════════════════════════════════════

def _cover_page(doc: Document, fd: FinancialData):
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(fd.company_name)
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = _NAVY

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Financial Analysis Report")
    r.font.size = Pt(16)
    r.font.color.rgb = _GOLD

    if fd.ticker:
        tk = doc.add_paragraph()
        tk.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = tk.add_run(f"{fd.ticker}  ·  {fd.exchange or 'NSE/BSE'}")
        r2.font.size = Pt(12)
        r2.font.color.rgb = _NAVY

    years_str = "  |  ".join(str(y) for y in fd.sorted_years())
    yr_para = doc.add_paragraph()
    yr_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    yr_r = yr_para.add_run(f"Fiscal Years: {years_str}")
    yr_r.font.size = Pt(12)

    doc.add_paragraph()
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_r = date_para.add_run(
        f"Generated: {datetime.now().strftime('%B %d, %Y')}"
    )
    date_r.font.size = Pt(10)
    date_r.font.italic = True

    for _ in range(6):
        doc.add_paragraph()

    disc = doc.add_paragraph(
        "This report is generated by Financial Analysis AI for informational purposes only. "
        "It does not constitute investment advice. All data sourced from company filings."
    )
    disc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in disc.runs:
        run.font.size = Pt(9)
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.add_page_break()


def _income_table(doc: Document, fd: FinancialData):
    years = fd.sorted_years()
    headers = ["Metric", "Unit"] + [str(y) for y in years]
    rows_data = []

    metrics = [
        ("Revenue", "₹ Cr", "revenue"),
        ("EBITDA",  "₹ Cr", "ebitda"),
        ("EBIT",    "₹ Cr", "ebit"),
        ("PAT / Net Income", "₹ Cr", "pat"),
        ("Gross Margin",    "%",    "gross_margin_pct"),
        ("EBITDA Margin",   "%",    "ebitda_margin_pct"),
        ("Net Margin",      "%",    "net_margin_pct"),
        ("EPS (Basic)",     "₹",   "eps_basic"),
        ("Dividends/Share", "₹",   "dividends_per_share"),
    ]

    for label, unit, field in metrics:
        row = [label, unit]
        for yr in years:
            s = fd.get_income_statement(yr)
            if s is None:
                row.append("—")
                continue
            if field == "gross_margin_pct":
                v = (s.gross_profit / s.revenue * 100) if s.gross_profit and s.revenue else None
                row.append(_fmt_num(v, 1) if v is not None else "—")
            elif field == "ebitda_margin_pct":
                v = (s.ebitda / s.revenue * 100) if s.ebitda and s.revenue else None
                row.append(_fmt_num(v, 1) if v is not None else "—")
            elif field == "net_margin_pct":
                v = (s.pat / s.revenue * 100) if s.pat and s.revenue else None
                row.append(_fmt_num(v, 1) if v is not None else "—")
            else:
                v = getattr(s, field, None)
                row.append(_fmt_num(v, 2 if unit == "₹" else 0))
        rows_data.append(row)

    _add_multi_col_table(doc, headers, rows_data)


def _balance_sheet_table(doc: Document, fd: FinancialData):
    years = fd.sorted_years()
    headers = ["Metric", "Unit"] + [str(y) for y in years]
    rows_data = []

    metrics = [
        ("Total Assets",           "₹ Cr", "total_assets"),
        ("Total Current Assets",   "₹ Cr", "total_current_assets"),
        ("Net Fixed Assets",       "₹ Cr", "net_fixed_assets"),
        ("Total Equity",           "₹ Cr", "total_equity"),
        ("Long-Term Debt",         "₹ Cr", "long_term_debt"),
        ("Short-Term Borrowings",  "₹ Cr", "short_term_borrowings"),
        ("Total Current Liabilities","₹ Cr","total_current_liabilities"),
        ("Cash & Equivalents",     "₹ Cr", "cash_and_equivalents"),
    ]

    for label, unit, field in metrics:
        row = [label, unit]
        for yr in years:
            s = fd.get_balance_sheet(yr)
            v = getattr(s, field, None) if s else None
            row.append(_fmt_num(v))
        rows_data.append(row)

    _add_multi_col_table(doc, headers, rows_data)


def _cash_flow_table(doc: Document, fd: FinancialData):
    years = fd.sorted_years()
    headers = ["Metric", "Unit"] + [str(y) for y in years]
    rows_data = []

    metrics = [
        ("Cash from Operations (CFO)", "₹ Cr", "cash_from_operations"),
        ("Cash from Investing  (CFI)", "₹ Cr", "cash_from_investing"),
        ("Cash from Financing  (CFF)", "₹ Cr", "cash_from_financing"),
        ("Capital Expenditure (Capex)","₹ Cr", "capex"),
        ("Free Cash Flow (FCF)",       "₹ Cr", "free_cash_flow"),
        ("Closing Cash Balance",       "₹ Cr", "closing_cash"),
    ]

    for label, unit, field in metrics:
        row = [label, unit]
        for yr in years:
            s = fd.get_cash_flow(yr)
            v = getattr(s, field, None) if s else None
            row.append(_fmt_num(v))
        rows_data.append(row)

    _add_multi_col_table(doc, headers, rows_data)


def _ratios_table(doc: Document, fd: FinancialData):
    years = fd.sorted_years()
    latest = years[-1]
    is_l = fd.get_income_statement(latest)
    bs_l = fd.get_balance_sheet(latest)
    cf_l = fd.get_cash_flow(latest)
    if not is_l or not bs_l:
        return

    def safe_div(a, b, pct=False, mult=100):
        if a is None or b is None or b == 0:
            return "—"
        v = a / b
        return f"{v * mult:.1f}%" if pct else f"{v:.2f}x"

    total_debt = (bs_l.long_term_debt or 0) + (bs_l.short_term_borrowings or 0)
    wc = (bs_l.total_current_assets or 0) - (bs_l.total_current_liabilities or 0)
    inv = bs_l.inventory or 0
    recv = bs_l.accounts_receivable or 0

    rows = [
        ("LIQUIDITY", ""),
        ("Current Ratio",   safe_div(bs_l.total_current_assets, bs_l.total_current_liabilities)),
        ("Quick Ratio",     safe_div((bs_l.total_current_assets or 0) - inv, bs_l.total_current_liabilities)),
        ("Cash Ratio",      safe_div((bs_l.cash_and_equivalents or 0) + (bs_l.short_term_investments or 0), bs_l.total_current_liabilities)),
        ("PROFITABILITY", ""),
        ("Gross Margin",    safe_div(is_l.gross_profit, is_l.revenue, pct=True)),
        ("EBITDA Margin",   safe_div(is_l.ebitda, is_l.revenue, pct=True)),
        ("EBIT Margin",     safe_div(is_l.ebit, is_l.revenue, pct=True)),
        ("Net Margin",      safe_div(is_l.pat, is_l.revenue, pct=True)),
        ("ROE",             safe_div(is_l.pat, bs_l.total_equity, pct=True)),
        ("ROA",             safe_div(is_l.pat, bs_l.total_assets, pct=True)),
        ("ROCE",            safe_div(is_l.ebit, (bs_l.total_assets or 0) - (bs_l.total_current_liabilities or 0), pct=True)),
        ("LEVERAGE", ""),
        ("Debt/Equity",         safe_div(total_debt, bs_l.total_equity)),
        ("Interest Coverage",   safe_div(is_l.ebit, is_l.interest_expense)),
        ("Net Debt (₹ Cr)",     _fmt_num(total_debt - (bs_l.cash_and_equivalents or 0) - (bs_l.short_term_investments or 0))),
        ("EFFICIENCY", ""),
        ("Asset Turnover",          safe_div(is_l.revenue, bs_l.total_assets)),
        ("DSO (days)",              _fmt_num((recv / is_l.revenue * 365) if is_l.revenue and recv else None, 0)),
        ("DIO (days)",              _fmt_num((inv / (is_l.cogs or 1) * 365) if inv and is_l.cogs else None, 0)),
        ("CASH FLOW", ""),
        ("FCF (₹ Cr)",             _fmt_num(cf_l.free_cash_flow if cf_l else None)),
        ("FCF / PAT",              safe_div(cf_l.free_cash_flow if cf_l else None, is_l.pat)),
        ("Capex / Revenue",        safe_div(abs(cf_l.capex) if cf_l and cf_l.capex else None, is_l.revenue, pct=True)),
    ]

    table = doc.add_table(rows=len(rows) + 1, cols=2)
    table.style = "Table Grid"

    hdr = table.rows[0]
    for cell, text in zip(hdr.cells, ["Ratio / Metric", f"FY{latest}"]):
        cell.text = text
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = _WHITE
        _set_cell_bg(cell, "1F3864")

    for i, (label, value) in enumerate(rows):
        row = table.rows[i + 1]
        is_section = value == ""
        row.cells[0].text = label
        row.cells[1].text = str(value)
        row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if is_section:
            _set_cell_bg(row.cells[0], "C9A84C")
            _set_cell_bg(row.cells[1], "C9A84C")
            row.cells[0].paragraphs[0].runs[0].font.bold = True
            row.cells[0].paragraphs[0].runs[0].font.color.rgb = _WHITE
        elif i % 2 == 1:
            _set_cell_bg(row.cells[0], "EBF3FB")
            _set_cell_bg(row.cells[1], "EBF3FB")

    for row in table.rows:
        row.cells[0].width = Inches(3.2)
        row.cells[1].width = Inches(1.5)

    doc.add_paragraph()


def _valuation_summary(doc: Document, fd: FinancialData):
    rows = [
        ("DCF Valuation — Intrinsic Value/Share", "See DCF sheet → B33"),
        ("DDM Valuation — Intrinsic Value/Share", "See DDM sheet → B24"),
        ("EV/EBITDA Comps — Implied Price",        "See Comps sheet → D24"),
        ("P/E Comps — Implied Price",              "See Comps sheet → D26"),
        ("Valuation Range (Low–High)",             "See Comps sheet → C35:C36"),
        ("Central Estimate",                       "See Comps sheet → C37"),
        ("Current Market Price",                   "Update Settings → B12"),
    ]
    _add_kv_table(doc, rows, col_widths=(3.5, 2.5))
    _body(doc,
          "To populate exact values, open the Excel workbook, enter the Market Cap and Current "
          "Share Price in the Settings sheet, and the valuation summary will auto-calculate.")


# ══════════════════════════════════════════════════════════════════════════════
# Main entry point
# ══════════════════════════════════════════════════════════════════════════════

def generate_report(
    financial_data: FinancialData,
    commentary: Dict[str, str],
    output_path: Optional[str] = None,
    excel_path: Optional[str] = None,
) -> str:
    """
    Generate a Word (.docx) report.
    Returns the path to the generated file.
    """
    if output_path is None:
        safe = financial_data.company_name.replace(" ", "_").replace("/", "-")
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_path = str(REPORTS_DIR / f"{safe}_Report_{ts}.docx")

    log.info(f"Generating Word report → {output_path}")
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    # Default paragraph font
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    # ── 1. Cover ─────────────────────────────────────────────────────────────
    _cover_page(doc, financial_data)

    # ── 2. Executive Summary ─────────────────────────────────────────────────
    _heading(doc, "1.  Executive Summary")
    _body(doc, commentary.get("executive_summary", ""))
    doc.add_paragraph()

    # ── 3. Revenue & Profitability ───────────────────────────────────────────
    _heading(doc, "2.  Revenue & Profitability Analysis")
    _body(doc, commentary.get("revenue_analysis", ""))
    doc.add_paragraph()
    _body(doc, commentary.get("profitability_analysis", ""))
    doc.add_paragraph()
    _heading(doc, "Income Statement Summary", level=2)
    _income_table(doc, financial_data)

    doc.add_page_break()

    # ── 4. Balance Sheet ─────────────────────────────────────────────────────
    _heading(doc, "3.  Balance Sheet Analysis")
    _body(doc, commentary.get("balance_sheet_analysis", ""))
    doc.add_paragraph()
    _heading(doc, "Balance Sheet Summary", level=2)
    _balance_sheet_table(doc, financial_data)

    # ── 5. Cash Flow ─────────────────────────────────────────────────────────
    _heading(doc, "4.  Cash Flow Analysis")
    _body(doc, commentary.get("cash_flow_analysis", ""))
    doc.add_paragraph()
    _heading(doc, "Cash Flow Summary", level=2)
    _cash_flow_table(doc, financial_data)

    doc.add_page_break()

    # ── 6. Key Ratios ────────────────────────────────────────────────────────
    _heading(doc, "5.  Key Ratios Dashboard")
    _ratios_table(doc, financial_data)

    # ── 7. Valuation Summary ─────────────────────────────────────────────────
    _heading(doc, "6.  Valuation Summary")
    _body(doc,
          "The following valuation methods have been applied. All figures update automatically "
          "once Market Cap and Share Price are entered in the Excel Settings sheet.")
    _valuation_summary(doc, financial_data)

    doc.add_page_break()

    # ── 8. Risks & Strengths ─────────────────────────────────────────────────
    _heading(doc, "7.  Key Risks & Strengths")

    _heading(doc, "Key Strengths", level=2)
    strengths = commentary.get("key_strengths", "")
    for bullet in strengths.split("\n"):
        b = bullet.strip()
        if b:
            p = doc.add_paragraph(b.lstrip("•").strip(), style="List Bullet")
            p.runs[0].font.size = Pt(11)

    doc.add_paragraph()
    _heading(doc, "Key Risks", level=2)
    risks = commentary.get("key_risks", "")
    for bullet in risks.split("\n"):
        b = bullet.strip()
        if b:
            p = doc.add_paragraph(b.lstrip("•").strip(), style="List Bullet")
            p.runs[0].font.size = Pt(11)

    # ── 9. Disclaimer ────────────────────────────────────────────────────────
    doc.add_page_break()
    _heading(doc, "Disclaimer", level=2)
    disc_text = (
        "This report has been generated by Financial Analysis AI and is intended solely for "
        "informational and educational purposes. It does not constitute investment advice, "
        "a solicitation to buy or sell any security, or a recommendation of any kind. "
        "All financial data has been sourced from publicly available company filings and may "
        "be subject to restatement. Forward-looking projections (DCF, DDM) are based on "
        "assumptions that may not materialise. Independent verification is strongly recommended "
        "before making any investment decision. Past performance is not indicative of future results."
    )
    _body(doc, disc_text)
    p = doc.add_paragraph()
    r = p.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y %H:%M')}")
    r.font.size = Pt(9)
    r.font.italic = True

    doc.save(output_path)
    log.info(f"Report saved: {output_path}")
    return output_path
